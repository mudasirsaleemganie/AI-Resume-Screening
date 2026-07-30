import io
import re
from typing import Iterable

from docx import Document
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from .taxonomy import ROLE_PROFILES, SKILL_ALIASES

MODEL_VERSION = "hybrid-tfidf-skills-v1.0"


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\x00", " ")).strip()


def parse_file(filename: str, content: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif lower.endswith(".docx"):
        document = Document(io.BytesIO(content))
        text = "\n".join(p.text for p in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                text += "\n" + " ".join(cell.text for cell in row.cells)
    else:
        raise ValueError("Only PDF and DOCX files are supported")
    text = normalize(text)
    if len(text) < 40:
        raise ValueError("The resume has too little extractable text; use a text-based PDF or DOCX")
    return text[:100_000]


def contains_term(text: str, term: str) -> bool:
    pattern = rf"(?<![\w+#]){re.escape(term.lower())}(?![\w+#])"
    return bool(re.search(pattern, text.lower()))


def extract_skills(text: str) -> list[str]:
    found = []
    for canonical, aliases in SKILL_ALIASES.items():
        if any(contains_term(text, alias) for alias in aliases):
            found.append(canonical)
    return found


def extract_contact(text: str) -> dict:
    email = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", text)
    phone = re.search(r"(?:\+?\d{1,3}[-.\s]?)?(?:\d[-.\s]?){10}", text)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    possible_name = lines[0][:80] if lines else ""
    if any(token in possible_name.lower() for token in ("resume", "curriculum", "@", "http")):
        possible_name = ""
    return {
        "name": possible_name,
        "email": email.group(0) if email else "",
        "phone": re.sub(r"\s+", " ", phone.group(0)).strip() if phone else "",
    }


def extract_experience(text: str) -> float:
    explicit = [
        float(value)
        for value in re.findall(r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?experience", text.lower())
    ]
    if explicit:
        return min(max(explicit), 50)
    years = [int(y) for y in re.findall(r"\b(?:19|20)\d{2}\b", text)]
    if len(years) >= 2:
        return float(min(max(years) - min(years), 50))
    return 0.0


def extract_education(text: str) -> list[str]:
    patterns = {
        "Master of Computer Applications (MCA)": r"\b(?:mca|master of computer applications?)\b",
        "Bachelor of Computer Applications (BCA)": r"\b(?:bca|bachelor of computer applications?)\b",
        "Master's degree": r"\b(?:master'?s|m\.?tech|m\.?sc)\b",
        "Bachelor's degree": r"\b(?:bachelor'?s|b\.?tech|b\.?e\.?|b\.?sc)\b",
        "Diploma": r"\bdiploma\b",
    }
    return [label for label, pattern in patterns.items() if re.search(pattern, text, re.I)]


def semantic_similarity(resume: str, job: str) -> float:
    vectors = TfidfVectorizer(stop_words="english", ngram_range=(1, 2), max_features=5000).fit_transform([resume, job])
    return float(cosine_similarity(vectors[0:1], vectors[1:2])[0][0] * 100)


def predict_role(resume: str) -> tuple[str, float]:
    names = list(ROLE_PROFILES)
    documents = [resume] + [ROLE_PROFILES[name] for name in names]
    matrix = TfidfVectorizer(stop_words="english", ngram_range=(1, 2)).fit_transform(documents)
    similarities = cosine_similarity(matrix[0:1], matrix[1:]).flatten()
    index = int(similarities.argmax())
    confidence = min(98.0, 45.0 + float(similarities[index]) * 100)
    return names[index], round(confidence, 1)


def canonicalize_required(skills: Iterable[str]) -> list[str]:
    values = []
    for raw in skills:
        lowered = raw.strip().lower()
        canonical = next(
            (name for name, aliases in SKILL_ALIASES.items() if lowered == name or lowered in aliases),
            lowered,
        )
        if canonical and canonical not in values:
            values.append(canonical)
    return values


def analyze(text: str, job_description: str, required_skills: list[str], minimum_experience: float) -> dict:
    skills = extract_skills(text)
    required = canonicalize_required(required_skills) or extract_skills(job_description)
    matched = [skill for skill in required if skill in skills]
    missing = [skill for skill in required if skill not in skills]
    skill_score = 100.0 if not required else len(matched) / len(required) * 100
    semantic_score = semantic_similarity(text, job_description)
    experience_years = extract_experience(text)
    experience_score = 100.0 if minimum_experience <= 0 else min(experience_years / minimum_experience * 100, 100)
    overall = skill_score * 0.55 + semantic_score * 0.30 + experience_score * 0.15
    role, confidence = predict_role(text)
    strengths = []
    if skill_score >= 70:
        strengths.append("Strong coverage of the job's required skills")
    if semantic_score >= 45:
        strengths.append("Resume content is well aligned with the job description")
    if experience_score >= 100 and minimum_experience > 0:
        strengths.append("Meets the stated experience requirement")
    if not strengths:
        strengths.append("Shows transferable experience that can be developed further")
    recommendations = []
    if missing:
        recommendations.append(f"Build evidence for these missing skills: {', '.join(missing[:6])}")
    if semantic_score < 45:
        recommendations.append("Add measurable achievements and job-relevant terminology naturally")
    if experience_score < 100:
        recommendations.append("Clarify employment and project dates so experience can be verified")
    recommendations.append("Keep the resume concise, truthful and tailored to each application")
    return {
        "skills": skills,
        "experienceYears": round(experience_years, 1),
        "analysis": {
            "overallScore": round(overall, 1),
            "skillScore": round(skill_score, 1),
            "semanticScore": round(semantic_score, 1),
            "experienceScore": round(experience_score, 1),
            "matchedSkills": matched,
            "missingSkills": missing,
            "predictedRole": role,
            "roleConfidence": confidence,
            "strengths": strengths,
            "recommendations": recommendations,
            "modelVersion": MODEL_VERSION,
        },
    }

